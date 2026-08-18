#!/usr/bin/env python3
"""Generate the CompTIA A+ (Core 1 and Core 2) Lesson Plan (LP) DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 5-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea counted within, final assessment Day 5 4:30pm). Topics/labs come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5; from data_domain6 import DOMAIN6
from data_domain7 import DOMAIN7; from data_domain8 import DOMAIN8
from data_domain9 import DOMAIN9
ACT=(DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5+DOMAIN6+DOMAIN7+DOMAIN8+DOMAIN9)
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","10:00",30,"admin","Welcome, trainer and learner introductions, ground rules, learning outcomes, course outline and mandatory digital attendance (AM)"),
    ("10:00","11:00",60,"topic","Topic 1 — Mobile Devices: laptop hardware and field-replaceable units, display technologies (LCD/IPS/VA/OLED), digitizers, connection methods, docking (concepts + demo)"),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:30",135,"lab","Hands-on: "+lab_titles([1,2,3,4])),
    ("13:30","14:30",60,"lunch","Lunch break"),
    ("14:30","16:00",90,"topic","Topic 2 — Networking: TCP/UDP, the A+ port list, network devices, wireless standards and bands, SOHO services (concepts + demo). Digital attendance (PM)"),
    ("16:00","16:15",15,"break","Tea break"),
    ("16:15","18:15",120,"lab","Hands-on: "+lab_titles([5,6,7])),
    ("18:15","18:30",15,"recap","Day 1 recap, Q&A and PM digital attendance"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Topic 2 continues — channel planning, SOHO configuration, cabling and packet analysis. Hands-on: "+lab_titles([8,9])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:30",135,"lab","Hands-on: "+lab_titles([10,11])+". Topic 3 — Hardware: cables and connectors, RAM, storage and RAID (concepts)"),
    ("13:30","14:30",60,"lunch","Lunch break"),
    ("14:30","16:00",90,"lab","Hands-on: "+lab_titles([12,13,14])+". Digital attendance (PM)"),
    ("16:00","16:15",15,"break","Tea break"),
    ("16:15","18:15",120,"lab","Topic 3 continues — motherboards, BIOS/UEFI, CPUs and cooling, power supplies, printers. Hands-on: "+lab_titles([15,16])),
    ("18:15","18:30",15,"recap","Day 2 recap, Q&A and PM digital attendance"),
 ]),
 3: (C.DAY_THEMES[3], [
    ("9:30","9:45",15,"recap","Day 2 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Hands-on: "+lab_titles([17,18])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:30",135,"topic","Topic 4 — Virtualization and Cloud Computing: hypervisor types, VM resource planning, cloud service and deployment models (concepts). Hands-on: "+lab_titles([19,20,21])),
    ("13:30","14:30",60,"lunch","Lunch break"),
    ("14:30","16:00",90,"topic","Topic 5 — Hardware and Network Troubleshooting: the CompTIA six-step methodology, POST and boot faults (concepts). Hands-on: "+lab_titles([22,23])+". Digital attendance (PM)"),
    ("16:00","16:15",15,"break","Tea break"),
    ("16:15","18:15",120,"lab","Hands-on: "+lab_titles([24,25,26])),
    ("18:15","18:30",15,"recap","Day 3 recap, Q&A and PM digital attendance"),
 ]),
 4: (C.DAY_THEMES[4], [
    ("9:30","9:45",15,"recap","Day 3 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Hands-on: "+lab_titles([27,28])+". End of Core 1 domains"),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:30",135,"topic","Topic 6 — Operating Systems: Windows editions and installation, partitioning, command line, Windows tools (concepts). Hands-on: "+lab_titles([29,30,31])),
    ("13:30","14:30",60,"lunch","Lunch break"),
    ("14:30","16:00",90,"lab","Hands-on: "+lab_titles([32,33,34])+". Digital attendance (PM)"),
    ("16:00","16:15",15,"break","Tea break"),
    ("16:15","18:15",120,"lab","Hands-on: "+lab_titles([35,36])+". Topic 7 — Security: physical and logical controls, authentication (concepts). Hands-on: "+lab_titles([37])),
    ("18:15","18:30",15,"recap","Day 4 recap, Q&A and PM digital attendance"),
 ]),
 5: (C.DAY_THEMES[5], [
    ("9:30","9:45",15,"recap","Day 4 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Topic 7 continues — malware, social engineering, wireless and workstation hardening. Hands-on: "+lab_titles([38,39,40])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","12:45",90,"lab","Hands-on: "+lab_titles([41,42,43,44])),
    ("12:45","13:30",45,"topic","Topic 8 — Software Troubleshooting: Windows symptoms and recovery tools, malware and browser symptoms (concepts). Hands-on: "+lab_titles([45])),
    ("13:30","14:30",60,"lunch","Lunch break"),
    ("14:30","15:45",75,"lab","Hands-on: "+lab_titles([46,47,48,49])+". Digital attendance (PM)"),
    ("15:45","16:15",30,"topic","Topic 9 — Operational Procedures: documentation, change and asset management, backup, safety, professionalism. Hands-on: "+lab_titles([50,51,52,53,54])),
    ("16:15","16:30",15,"assess","Course recap, exam preparation guidance, TRAQOM survey and Briefing for Assessment"),
    ("16:30","17:30",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book"),
    ("17:30","18:30",60,"assess","Practical Performance (PP) — hands-on support, configuration and troubleshooting tasks, 1 hour, open book. Assessment digital attendance"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("6.0","1 March 2025","Previous release — 5-day lesson plan covering the nine A+ Core 1 and Core 2 topics.",C.TRAINER),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  f"Major revision — rebuilt against the current CompTIA A+ Core 1 (220-1101) and Core 2 (220-1102) exam domain weightings. "
  f"Schedule now covers {len(ACT)} hands-on labs across the five days, using the browser-based lab toolkit "
  f"(IP Calculator, PCAP Analyzer, Cybersecurity Simulator, RegexLab) and the Killercoda Ubuntu playground.",C.TRAINER),
 ("7.1",C.VERSION_DATE,"Aligned the Skills Framework section to the full competency set the assessment tests — K1–K6 (Written Assessment) and A1–A8 (Practical Performance) — so no learner is assessed on an outcome the courseware did not declare.",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("TSC Alignment",f"{C.TSC_TITLE}  ({C.TSC_CODE})"),
      ("Duration",f"{C.DAYS} days · 8 training hours per day ({C.HOURS} hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks counted within training time)"),
      ("Mode","Instructor-led, hands-on IT support labs using physical hardware where available and browser-based tools throughout"),
      ("Certification Alignment","CompTIA A+ Core 1 (220-1101) and Core 2 (220-1102)"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Skills Framework Alignment",1)
doc.add_paragraph(f"This course is aligned to the Skills Framework Technical Skill and Competency "
                  f"{C.TSC_TITLE} ({C.TSC_CODE}).")
doc.add_paragraph("Knowledge:")
for k in C.TSC_KNOWLEDGE:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(k).font.size=Pt(10)
doc.add_paragraph("Abilities:")
for a in C.TSC_ABILITIES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          f"Final assessment is conducted on Day {C.DAYS} from 4:30 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

H("Lab Toolkit",1)
doc.add_paragraph("All hands-on labs use free, browser-based tools that require no local installation, "
                  "so every learner works in an identical environment:")
tk=doc.add_table(rows=0,cols=3); tk.style="Table Grid"
def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)
hdr=tk.add_row().cells
for i,htext in enumerate(["Tool","Link","Used for"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for name,url,desc in C.LAB_TOOLS:
    cells=tk.add_row().cells
    set_cell(cells[0],name,bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],url,size=9)
    set_cell(cells[2],desc,size=9)
for row in tk.rows:
    row.cells[0].width=Inches(1.5); row.cells[1].width=Inches(2.2); row.cells[2].width=Inches(3.1)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to CompTIA A+ exam domains)",1)
tt=doc.add_table(rows=0,cols=5); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic / Exam domain","Exam","Exam weighting","Course time","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9,fill=TOPIC_FILL)
    set_cell(cells[1],tp["core"],size=9,fill=TOPIC_FILL)
    set_cell(cells[2],tp["exam_weight"],size=9,fill=TOPIC_FILL)
    set_cell(cells[3],tp["weighting"],size=9,fill=TOPIC_FILL)
    set_cell(cells[4],", ".join(f"Lab {a['num']}" for a in acts),size=9)
for row in tt.rows:
    row.cells[0].width=Inches(2.3); row.cells[1].width=Inches(0.7)
    row.cells[2].width=Inches(0.9); row.cells[3].width=Inches(0.7); row.cells[4].width=Inches(2.2)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
